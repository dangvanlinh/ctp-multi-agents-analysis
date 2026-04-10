"""Generate VIP Level System Design Doc v2 — Word document for dev."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = 'Arial'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        p = c.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = c._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '1E2761'
        })
        shading.append(bg)
    for r, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if r % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F5'
                })
                shading.append(bg)
    doc.add_paragraph()
    return t


def decision(title, chosen, pro, con):
    """Decision box: chỉ chốt + pro + con, không list options."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f'Edge Case: {title}')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('Chốt: ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x2E, 0xA0, 0x6A)
    r2 = p.add_run(chosen)
    r2.font.size = Pt(10)
    r2.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('Pro: ')
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(pro).font.size = Pt(9)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('Con: ')
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(con).font.size = Pt(9)


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.3)
    p.runs[0].font.size = Pt(10)


def num_item(text):
    p = doc.add_paragraph(text, style='List Number')
    p.runs[0].font.size = Pt(10)


def red_text(text):
    """Text đỏ cho phần chưa chốt."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xE6, 0x3E, 0x31)
    r.font.size = Pt(11)


def add_screenshot(path, width=Inches(5.5)):
    """Insert screenshot image."""
    try:
        doc.add_picture(path, width=width)
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        p = doc.add_paragraph(f'[Screenshot: {path} - {e}]')
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ═══════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('VIP LEVEL SYSTEM')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)

p = doc.add_paragraph()
r = p.add_run('Design Document for Development')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
r = p.add_run('Status: Draft | Date: 2026-04-06 | UI Mockup: vip_ui_mockup.html')
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_paragraph()

# ═══ 1. OVERVIEW ═══
doc.add_heading('1. Overview', level=1)
doc.add_paragraph('Redesign VIP t\u1EEB subscription \u0111\u01A1n gi\u1EA3n (100k/10 ng\u00E0y, 82.8% mua 1 l\u1EA7n) th\u00E0nh h\u1EC7 th\u1ED1ng level progression v\u1EDBi loss aversion.')
p = doc.add_paragraph()
p.add_run('Core Pillars:').bold = True
num_item('Fix 2 v\u1EA5n \u0111\u1EC1 VIP: repay trong th\u00E1ng (ARPT 1\u21923) + repay th\u00E1ng sau (repeat 17%\u219250%)')
num_item('Kh\u00F4ng c\u1EAFn nhau v\u1EDBi event \u0111ang ch\u1EA1y (SKB, CH\u0110, Tr\u1EE9ng, Lazer)')

# ═══ 2. KPI TARGET (moved up from 12) ═══
doc.add_heading('2. KPI Target', level=1)
p = doc.add_paragraph()
r = p.add_run('Rev target: 350tr+/th\u00E1ng (hi\u1EC7n ~72tr)')
r.bold = True; r.font.size = Pt(13)
p = doc.add_paragraph()
r = p.add_run('Model: 800 PU m\u1EDBi, new retain 50%, old retain 80%, new ARPT=2, old ARPT=3')
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
add_table(
    ['Th\u00E1ng', 'PU m\u1EDBi', 'Old retained', 'Active PU', 'L\u01B0\u1EE3t', 'Rev'],
    [
        ['T1', '800', '362', '1,162', '2,686', '269M'],
        ['T2', '800', '690', '1,490', '3,670', '367M'],
        ['T3', '800', '952', '1,752', '4,456', '446M'],
        ['T4', '800', '1,162', '1,962', '5,086', '509M'],
        ['T6+', '800', '~2,000', '~2,800', '~7,600', '~760M'],
    ])

# ═══ 3. CORE MECHANIC ═══
doc.add_heading('3. Core Mechanic', level=1)
add_table(['Rule', 'Chi ti\u1EBFt'], [
    ['S\u1ED1 c\u1EA5p', '5 c\u1EA5p (Lv1-Lv5)'],
    ['Gi\u00E1', '100k VN\u0110 / 10 ng\u00E0y (t\u1EA5t c\u1EA3 level)'],
    ['Level up', 'Mua li\u00EAn ti\u1EBFp = +1 Level'],
    ['Reset', 'Kh\u00F4ng mua + h\u1EBFt 3 ng\u00E0y grace = Reset v\u1EC1 Lv0'],
    ['Hi\u1EC3n th\u1ECB', 'Ch\u1EC9 th\u1EA5y benefit level ti\u1EBFp theo +1, level xa b\u1ECB \u1EA9n'],
])

# ═══ 4. CONFIG BENEFIT ═══
doc.add_heading('4. Config Benefit', level=1)
p = doc.add_paragraph()
r = p.add_run('Gi\u00E1: 100k/10 ng\u00E0y. Benefit t\u00EDch l\u0169y c\u1ED9ng d\u1ED3n.')
r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
add_table(
    ['Lv', 'GEM', 'R.V\u00E0ng', 'Ch\u00ECa', 'R.T\u00EDm', 'EXP', 'R.Cam', 'Cosmetic', 'DKXX', 'Upg R', 'Bonus%G'],
    [
        ['1', '1,500G', '50', '10', '5', '\u2014', '5', 'Khung+FX Vip1', '\u2014', '\u2014', 'TBD'],
        ['2', '1,500G', '50', '10', '10', 'x2', '10', 'Khung+FX Vip2', '+3%', '\u2014', 'TBD'],
        ['3', '1,500G', '50', '10', '15', 'x3', '15', 'Khung+FX Vip3', '+5%', 'Unlock', 'TBD'],
        ['4', '1,500G', '50', '10', '20', 'x3', '15', 'Khung+FX Vip4', '+8%', '+10%', 'TBD'],
        ['5', '1,500G', '50', '10', '30', 'x4', '20', 'Frame lobby', '+10%', '+15%', 'TBD'],
    ])

# ═══ 5. BENEFIT DELIVERY ═══
doc.add_heading('5. Benefit Delivery \u2014 2 Lo\u1EA1i', level=1)

doc.add_heading('5.1 Daily Claim (ph\u1EA3i login nh\u1EADn)', level=2)
doc.add_paragraph('Kh\u00F4ng login = m\u1EA5t qu\u00E0 ng\u00E0y \u0111\u00F3. K\u00EDch new user login h\u00E0ng ng\u00E0y \u2192 t\u0103ng retention.')
bullet('GEM (83G/ng\u00E0y)')
bullet('R\u01B0\u01A1ng v\u00E0ng (5/ng\u00E0y)')
bullet('R\u01B0\u01A1ng t\u00EDm: chia \u0111\u1EC1u 10 ng\u00E0y')
bullet('R\u01B0\u01A1ng cam: chia cho 5 ng\u00E0y cu\u1ED1i (ng\u00E0y 6-10)')
bullet('Ch\u00ECa kh\u00F3a: chia \u0111\u1EC1u 10 ng\u00E0y')
p = doc.add_paragraph()
p.add_run('VD Lv2 (10 t\u00EDm, 5 cam):').bold = True
add_table(['Ng\u00E0y', 'R.T\u00EDm', 'R.Cam'], [['1-5', '1/ng\u00E0y', '\u2014'], ['6-10', '1/ng\u00E0y', '1/ng\u00E0y']])

doc.add_heading('5.2 Auto-Active (mua VIP l\u00E0 c\u00F3)', level=2)
doc.add_paragraph('Whale c\u1EA7n buff khi ch\u01A1i, kh\u00F4ng c\u1EA7n login claim.')
bullet('x2/x3/x4 EXP')
bullet('DKXX boost')
bullet('Upgrade R access')
bullet('Cosmetic (khung, FX)')

decision('Benefit delivery: chia 2 lo\u1EA1i',
    'Chia 2 segment: Daily Claim (new) + Auto-Active (whale)',
    'New login h\u00E0ng ng\u00E0y, whale kh\u00F4ng b\u1ECB phi\u1EC1n',
    'Ph\u1EE9c t\u1EA1p h\u01A1n cho dev (2 lo\u1EA1i delivery)')

# ═══ 6. TIMING ═══
doc.add_heading('6. Timing & Countdown', level=1)

doc.add_heading('6.1 VIP t\u00EDnh theo ng\u00E0y', level=2)
bullet('Mua ng\u00E0y n\u00E0o th\u00EC h\u1EBFt 00:00 ng\u00E0y th\u1EE9 11')
bullet('Daily claim reset l\u00FAc 00:00')
decision('Timezone',
    'T\u00EDnh theo ng\u00E0y, reset 00:00',
    'D\u1EC5 hi\u1EC3u, align daily claim',
    'User mua 23:59 \u0111\u01B0\u1EE3c g\u1EA7n 1 ng\u00E0y \u2014 negligible')

doc.add_heading('6.2 Grace period', level=2)
bullet('H\u1EBFt 10 ng\u00E0y \u2192 user c\u00F3 3 ng\u00E0y \u0111\u1EC3 mua ti\u1EBFp')
bullet('Trong 3 ng\u00E0y: kh\u00F4ng qu\u00E0 daily, kh\u00F4ng buff auto-active')
bullet('Kh\u00F4ng mua trong 3 ng\u00E0y \u2192 reset Lv0')
decision('Reset level',
    'Reset Lv0 (m\u1EA5t h\u1EBFt)',
    'Loss aversion m\u1EA1nh nh\u1EA5t. 3 ng\u00E0y grace \u0111\u1EE7 d\u00E0i',
    'Harsh \u2014 nh\u01B0ng \u0111\u00F3 l\u00E0 point c\u1EE7a design')

# ═══ 7. MUA VIP ═══
doc.add_heading('7. Mua VIP \u2014 Rules', level=1)

doc.add_heading('7.1 Mua tr\u01B0\u1EDBc h\u1EA1n', level=2)
doc.add_paragraph('Khi user mua VIP ti\u1EBFp trong khi VIP hi\u1EC7n t\u1EA1i ch\u01B0a h\u1EBFt:')
num_item('Nh\u1EADn tr\u1ECDn qu\u00E0 c\u00F2n l\u1EA1i c\u1EE7a level hi\u1EC7n t\u1EA1i v\u00E0o inbox ngay')
num_item('Level hi\u1EC7n t\u1EA1i k\u1EBFt th\u00FAc')
num_item('Level m\u1EDBi k\u00EDch ho\u1EA1t ngay, 10 ng\u00E0y m\u1EDBi b\u1EAFt \u0111\u1EA7u')
decision('Mua tr\u01B0\u1EDBc h\u1EA1n',
    'Nh\u1EADn h\u1EBFt qu\u00E0 c\u00F2n l\u1EA1i + l\u00EAn level ngay',
    'User kh\u00F4ng m\u1EA5t g\u00EC, k\u00EDch mua s\u1EDBm = t\u0103ng ARPT',
    'Burst qu\u00E0 g\u1ED9p \u2014 nh\u01B0ng l\u00E0 qu\u00E0 \u0111\u00E1ng l\u1EBD nh\u1EADn r\u1ED3i')

doc.add_heading('7.2 Mua nhi\u1EC1u l\u1EA7n li\u1EC1n', level=2)
bullet('Kh\u00F4ng cho ph\u00E9p mua nhi\u1EC1u l\u1EA7n c\u00F9ng l\u00FAc \u0111\u1EC3 nh\u1EA3y level')
bullet('B\u1EAFt bu\u1ED9c 1 l\u1EA7n mua = 1 level')
bullet('Nh\u1EA3y level ch\u1EC9 d\u00E0nh cho offer new user / lapsed')
decision('Mua nhi\u1EC1u l\u1EA7n',
    'B\u1EAFt bu\u1ED9c mua t\u1EEBng level',
    'User tr\u1EA3i nghi\u1EC7m t\u1EEBng level, loss aversion m\u1EA1nh',
    'Whale kh\u00F4ng fast-track \u2014 nh\u01B0ng \u0111\u00F3 l\u00E0 intent')

# ═══ 8. NÚT KÍCH HOẠT ═══
doc.add_heading('8. N\u00FAt K\u00EDch Ho\u1EA1t \u2014 State Machine', level=1)
p = doc.add_paragraph()
r = p.add_run('Rule: Ch\u01B0a claim qu\u00E0 h\u00F4m nay \u2192 ch\u01B0a cho renew.')
r.bold = True; r.font.color.rgb = RGBColor(0xE6, 0x3E, 0x31)
add_table(['Tr\u1EA1ng th\u00E1i', 'N\u00FAt Claim', 'N\u00FAt K\u00EDch ho\u1EA1t'], [
    ['Ch\u01B0a l\u00E0 VIP', 'Kh\u00F4ng c\u00F3', 'Hi\u1EC7n (mua l\u1EA7n \u0111\u1EA7u)'],
    ['\u0110ang VIP, >3 ng\u00E0y, ch\u01B0a claim', 'Hi\u1EC7n', '\u1EA8n'],
    ['\u0110ang VIP, >3 ng\u00E0y, \u0111\u00E3 claim', 'Kh\u00F4ng', '\u1EA8n'],
    ['\u0110ang VIP, \u22643 ng\u00E0y, ch\u01B0a claim', 'Hi\u1EC7n', '\u1EA8n (ch\u1EDD claim)'],
    ['\u0110ang VIP, \u22643 ng\u00E0y, \u0111\u00E3 claim', 'Kh\u00F4ng', 'Hi\u1EC7n (renew)'],
    ['H\u1EBFt VIP, trong 3 ng\u00E0y grace', 'Kh\u00F4ng', 'Hi\u1EC7n (urgent)'],
    ['H\u1EBFt VIP, qu\u00E1 3 ng\u00E0y \u2192 Lv0', 'Kh\u00F4ng', 'Hi\u1EC7n (mua l\u1EA1i Lv1)'],
])

# ═══ 9. UPGRADE R ═══
doc.add_heading('9. Upgrade R', level=1)
add_table(['Rule', 'Chi ti\u1EBFt'], [
    ['Hi\u1EC7n t\u1EA1i', 'Ch\u1EC9 CLB h\u1EA1ng B+ m\u1EDBi upgrade R'],
    ['M\u1EDBi', 'VIP Lv3+ c\u0169ng unlock (m\u1EDF th\u00EAm \u0111\u01B0\u1EDDng, kh\u00F4ng gating m\u1EDBi)'],
    ['VIP h\u1EBFt', 'Kh\u00F3a ngay quy\u1EC1n Upgrade R'],
    ['Th\u1EBB R \u0111\u00E3 upgrade', 'Gi\u1EEF nguy\u00EAn, kh\u00F4ng revert'],
    ['CLB B + VIP', 'Bonus rate c\u1ED9ng d\u1ED3n'],
])
add_table(['Level', 'Upgrade R', 'M\u00F4 t\u1EA3'], [
    ['Lv1-2', '\u2014', 'Kh\u00F4ng c\u00F3'],
    ['Lv3', 'Unlock', 'M\u1EDF quy\u1EC1n upgrade R'],
    ['Lv4', '+10% rate', 'T\u0103ng t\u1EF7 l\u1EC7 th\u00E0nh c\u00F4ng'],
    ['Lv5', '+15% rate', 'T\u1EF7 l\u1EC7 cao nh\u1EA5t'],
])
decision('Kh\u00F3a khi VIP h\u1EBFt',
    'Kh\u00F3a ngay (instant action, kh\u00F4ng process d\u1EDF dang)',
    '\u0110\u01A1n gi\u1EA3n, t\u1EA1o urgency gi\u1EEF VIP active',
    'Kh\u00F4ng c\u00F3')
decision('CLB B + VIP c\u00F9ng l\u00FAc',
    'C\u1ED9ng d\u1ED3n bonus rate',
    'Kh\u00F4ng punish user \u0111\u00E3 \u0111\u1EA1t CLB B',
    'Rate upgrade R c\u00F3 th\u1EC3 cao \u2014 c\u1EA7n balance')

# ═══ 10. NOTI ═══
doc.add_heading('10. Noti/Communication', level=1)
bullet('Kh\u00F4ng push noti \u2014 d\u00F9ng in-game UI khi login')
bullet('Daily claim popup = touchpoint t\u1EF1 nhi\u00EAn \u0111\u1EC3 nh\u1EAFc countdown VIP')
bullet('User kh\u00F4ng login \u2192 kh\u00F4ng bi\u1EBFt VIP s\u1EAFp h\u1EBFt')

# ═══ 11. UI SCREENS ═══
doc.add_heading('11. UI Screens', level=1)
doc.add_paragraph('File mockup: vip_ui_mockup.html').italic = True

doc.add_heading('Screen 1: Ch\u01B0a VIP', level=2)
doc.add_paragraph('Preview Lv1, n\u00FAt K\u00EDch ho\u1EA1t, benefit m\u1EDD.')
add_screenshot('data/screenshots/screen_1.png')

doc.add_heading('Screen 2: \u0110ang VIP (Lv2)', level=2)
doc.add_paragraph('Benefit s\u00E1ng, n\u00FAt Claim qu\u00E0, countdown, arrows xem level.')
add_screenshot('data/screenshots/screen_2.png')

doc.add_heading('Screen 3: VIP \u22643 ng\u00E0y', level=2)
doc.add_paragraph('Urgent, n\u00FAt Claim + Renew (disabled \u0111\u1EBFn khi claim).')
add_screenshot('data/screenshots/screen_3.png')

doc.add_heading('Screen 4: H\u1EBFt VIP \u2014 3 ng\u00E0y grace', level=2)
doc.add_paragraph('T\u1EA5t c\u1EA3 benefit kh\u00F3a, countdown reset, n\u00FAt K\u00EDch ho\u1EA1t urgent.')
add_screenshot('data/screenshots/screen_4.png')

doc.add_heading('Arrow Navigation', level=2)
bullet('Tr\u00E1i: xem l\u1EA1i level \u0111\u00E3 qua')
bullet('Ph\u1EA3i: ch\u1EC9 t\u1EDBi level hi\u1EC7n t\u1EA1i +1, ti\u1EBFp theo \u201C???\u201D')
bullet('Lv5 ph\u1EA3i: \u201CMax Level\u201D')

# ═══ 12. CHƯA CHỐT (đỏ) ═══
doc.add_heading('12. Ch\u01B0a Ch\u1ED1t \u2014 C\u1EA7n Design Th\u00EAm', level=1)

red_text('\u26A0\uFE0F 12.1 Bonus %G (PARKED)')
bullet('Shop trung hi\u1EC7n cho bonus 5-20% free theo g\u00F3i (50k\u21925%, 1M\u219220%)')
bullet('65% GEM payer ch\u01B0a mua VIP \u2192 n\u1EBFu remove, h\u1ECD m\u1EA5t h\u1EBFt')
bullet('81% transactions l\u00E0 g\u00F3i \u2264200k (bonus \u226410%)')
bullet('H\u01B0\u1EDBng xem x\u00E9t: Gi\u1EEF shop 10% flat, VIP +5%/level on top')
bullet('Blocker: C\u1EA7n t\u00EDnh balance k\u1EF9, concern GEM inflation')

red_text('\u26A0\uFE0F 12.2 Lapsed Rule (PARKED)')
bullet('Offer skip Lv2 c\u00F3 th\u1EC3 b\u1ECB gaming (b\u1ECF VIP \u2192 ch\u1EDD KM \u2192 loop)')
bullet('Ch\u1EDD data live r\u1ED3i design')
bullet('Default: lapsed mua l\u1EA1i = Lv1, kh\u00F4ng KM')

red_text('\u26A0\uFE0F 12.3 New User Offer (\u0110\u00C3 CH\u1ED0T \u2014 c\u1EA7n UI)')
bullet('L\u1EA7n mua \u0111\u1EA7u ti\u00EAn trong \u0111\u1EDDi \u2192 skip l\u00EAn Lv2 ngay (gi\u00E1 100k)')
bullet('C\u1EA7n design offer UI mockup')

red_text('\u26A0\uFE0F 12.4 Art & UI Tasks')
bullet('Redesign n\u00FAt VIP \u1EDF bottom bar main screen')
bullet('New user offer popup')
bullet('Ref art th\u1EBB VIP c\u00E1c c\u1EA5p (Lv1-5)')
bullet('Ref khung qu\u00E0 trong UI VIP')
bullet('Ref cosmetic: khung avatar, FX tung XX, khung deco NV lobby, BG lobby')

out = 'data/VIP_Design_Doc_v4.docx'
doc.save(out)
print(f'Saved: {out}')
